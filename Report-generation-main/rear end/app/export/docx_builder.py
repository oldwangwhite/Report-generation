from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.entity.template import ReportTemplate
from app.utils.report_labels import report_type_label


class DocxBuilder:
    def build(
        self,
        report,
        outline: list,
        contents_by_chapter: dict[int, object],
        path: Path,
        template: ReportTemplate | None = None,
    ) -> None:
        document = self._new_document(template)
        structure = template.structure if template and template.structure else {}
        title_style = structure.get("titleStyle") or "Title"
        heading_style = structure.get("headingStyle") or "Heading 1"
        body_style = structure.get("bodyStyle") or "Normal"
        table_style = structure.get("tableStyle") or "Table Grid"

        self._setup_document(document)
        self._add_title(document, report, title_style)
        body_style_name = self._existing_style(document, body_style, "Normal")
        self._add_report_info_table(document, report, table_style)

        for chapter in outline:
            self._add_heading(document, chapter, heading_style)
            content = contents_by_chapter.get(chapter.id)
            if content and content.content:
                self._add_content_paragraphs(document, str(content.content), body_style_name)
            else:
                document.add_paragraph("本章节暂无保存内容。", style=body_style_name)
            for table_data in (content.tables if content else []) or []:
                self._add_table(document, table_data, table_style)

        document.save(path)

    def _new_document(self, template: ReportTemplate | None) -> DocxDocument:
        template_path = Path(template.file_path) if template and template.file_path else None
        if template_path and template_path.exists() and template_path.suffix.lower() == ".docx":
            document = Document(template_path)
            self._clear_body(document)
            return document
        return Document()

    def _setup_document(self, document: DocxDocument) -> None:
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        normal = document.styles["Normal"]
        normal.font.name = "宋体"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal.font.size = Pt(11)

    def _add_title(self, document: DocxDocument, report, title_style: Any) -> None:
        title = document.add_paragraph(style=self._existing_style(document, title_style, "Title"))
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(report.report_name)
        title_run.bold = True
        title_run.font.size = Pt(20)
        title_run.font.name = "黑体"
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    def _add_report_info_table(self, document: DocxDocument, report, table_style: Any) -> None:
        rows = [
            ("报告类型", report_type_label(report.report_type), "报告主题", report.topic or ""),
            ("电厂", report.plant or "", "专业", report.major or ""),
            ("年份", str(report.year or ""), "报告状态", report.status or ""),
        ]
        table = document.add_table(rows=len(rows), cols=4)
        table.style = self._existing_table_style(document, table_style, "Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_index, row_data in enumerate(rows):
            for col_index, value in enumerate(row_data):
                cell = table.rows[row_index].cells[col_index]
                cell.text = str(value)
                self._set_cell_margin(cell)
                if col_index in {0, 2}:
                    self._shade_cell(cell, "F2F6FC")
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        document.add_paragraph("")

    def _add_heading(self, document: DocxDocument, chapter, heading_style: Any) -> None:
        level = min(max(chapter.level, 1), 4)
        text = f"{chapter.chapter_no} {chapter.title}"
        if level == 1:
            heading = document.add_paragraph(style=self._existing_style(document, heading_style, "Heading 1"))
            run = heading.add_run(text)
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(31, 78, 121)
        else:
            heading = document.add_heading(text, level=level)
            for run in heading.runs:
                run.font.name = "黑体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    def _add_content_paragraphs(self, document: DocxDocument, content: str, body_style_name: str) -> None:
        parts = [part.strip() for part in content.replace("\r\n", "\n").split("\n") if part.strip()]
        if not parts:
            parts = [content.strip()]
        for part in parts:
            paragraph = document.add_paragraph(part, style=body_style_name)
            paragraph.paragraph_format.first_line_indent = Pt(22)
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _add_table(self, document: DocxDocument, table_data: dict, table_style: Any) -> None:
        title = document.add_paragraph(str(table_data.get("title") or "检查情况表"))
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].bold = True
        headers = [str(item) for item in table_data.get("headers", [])] or ["内容"]
        rows = table_data.get("rows", []) or []
        table = document.add_table(rows=1, cols=len(headers))
        table.style = self._existing_table_style(document, table_style, "Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = header
            self._shade_cell(cell, "D9EAF7")
            self._set_cell_margin(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True

        for row_data in rows:
            row = table.add_row()
            values = list(row_data)[: len(headers)]
            values.extend([""] * (len(headers) - len(values)))
            for index, value in enumerate(values):
                cell = row.cells[index]
                cell.text = str(value)
                self._set_cell_margin(cell)

        document.add_paragraph("")

    def _clear_body(self, document: DocxDocument) -> None:
        body = document.element.body
        for child in list(body):
            if child.tag.endswith("}sectPr"):
                continue
            body.remove(child)

    def _existing_style(self, document: DocxDocument, preferred: Any, fallback: str) -> str:
        style_name = str(preferred or fallback)
        try:
            document.styles[style_name]
            return style_name
        except KeyError:
            return fallback

    def _existing_table_style(self, document: DocxDocument, preferred: Any, fallback: str) -> str:
        style_name = str(preferred or fallback)
        try:
            document.styles[style_name]
            return style_name
        except KeyError:
            return fallback

    def _shade_cell(self, cell, color: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        tc_pr.append(shading)

    def _set_cell_margin(self, cell) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        margins = tc_pr.first_child_found_in("w:tcMar")
        if margins is None:
            margins = OxmlElement("w:tcMar")
            tc_pr.append(margins)
        for edge in ("top", "left", "bottom", "right"):
            tag = f"w:{edge}"
            element = margins.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                margins.append(element)
            element.set(qn("w:w"), "80")
            element.set(qn("w:type"), "dxa")
