from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from tests.test_content_stream import prepared_report_with_outline


def prepare_generated_report(client, headers):
    report_id, outline = prepared_report_with_outline(client, headers)
    chapter_id = outline[0]["chapterId"]
    client.put(
        f"/api/reports/{report_id}/chapters/{chapter_id}/content",
        headers=headers,
        json={
            "content": "用户编辑后的章节正文。",
            "tables": [
                {
                    "tableId": "tbl_001",
                    "title": "设备检查情况表",
                    "headers": ["序号", "设备", "问题", "整改建议"],
                    "rows": [["1", "主变压器", "温度偏高", "加强巡检"]],
                }
            ],
            "manualEdited": True,
        },
    )
    return report_id


def docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [p.text for p in document.paragraphs]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs + table_text)


def make_template_bytes() -> bytes:
    template = Document()
    if "CustomBody" not in [style.name for style in template.styles]:
        custom = template.styles.add_style("CustomBody", WD_STYLE_TYPE.PARAGRAPH)
        custom.base_style = template.styles["Normal"]
    template.add_paragraph("模板占位内容，导出时应被清空。")
    buffer = BytesIO()
    template.save(buffer)
    return buffer.getvalue()


def test_create_status_history_and_download_docx_uses_template(client, auth_headers, admin_headers):
    uploaded = client.post(
        "/api/templates",
        headers=admin_headers,
        data={"templateName": "自定义导出模板", "reportType": "summerCheck"},
        files={"file": ("template.docx", make_template_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    ).json()
    assert uploaded["code"] == 200
    template_id = uploaded["data"]["templateId"]
    client.put(
        f"/api/templates/{template_id}",
        headers=admin_headers,
        json={
            "templateName": "自定义导出模板",
            "status": "enabled",
            "structure": {
                "titleStyle": "Title",
                "headingStyle": "Heading 1",
                "bodyStyle": "CustomBody",
                "tableStyle": "Table Grid",
            },
        },
    )
    report_id = prepare_generated_report(client, auth_headers)

    created = client.post(
        f"/api/reports/{report_id}/exports",
        headers=auth_headers,
        json={"templateId": template_id, "fileFormat": "docx", "useLatestSavedContent": True},
    ).json()

    assert created["code"] == 200
    assert created["data"]["exportId"].startswith("exp_")
    assert created["data"]["status"] == "exporting"
    export_id = created["data"]["exportId"]

    status = client.get(
        f"/api/reports/{report_id}/exports/{export_id}", headers=auth_headers
    ).json()
    assert status["code"] == 200
    assert status["data"]["status"] == "exported"
    assert status["data"]["fileName"].endswith(".docx")
    assert status["data"]["downloadUrl"].endswith(f"/exports/{export_id}/download")

    history = client.get(
        f"/api/reports/{report_id}/exports",
        headers=auth_headers,
        params={"page": 1, "size": 10, "fileFormat": "docx"},
    ).json()
    assert history["code"] == 200
    assert history["data"]["total"] == 1
    assert history["data"]["items"][0]["exportId"] == export_id

    download = client.get(
        f"/api/reports/{report_id}/exports/{export_id}/download", headers=auth_headers
    )
    assert download.status_code == 200
    assert download.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "filename*=UTF-8''" in download.headers["content-disposition"]
    text = docx_text(download.content)
    assert "用户编辑后的章节正文。" in text
    assert "设备检查情况表" in text
    assert "报告类型：迎峰度夏检查报告" in text
    assert "summerCheck" not in text
    assert "模板占位内容" not in text

    document = Document(BytesIO(download.content))
    body_paragraph = next(p for p in document.paragraphs if p.text == "用户编辑后的章节正文。")
    assert body_paragraph.style.name == "CustomBody"


def test_pdf_export_returns_friendly_failure(client, auth_headers):
    report_id = prepare_generated_report(client, auth_headers)

    created = client.post(
        f"/api/reports/{report_id}/exports",
        headers=auth_headers,
        json={"fileFormat": "pdf", "useLatestSavedContent": True},
    ).json()

    assert created["code"] == 500
    assert created["message"] == "文件导出失败"
    assert created["data"]["errorType"] == "export_failed"