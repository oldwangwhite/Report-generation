from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.entity.template import ReportTemplate
from app.utils.report_labels import report_type_label


class DocxBuilder:
    def build(
        self,
        report,
        outline: list,
        contents_by_chapter: dict[int, object],
        path: Path,
        template: ReportTemplate | None = None,
    ) -> None:
        document = self._new_document(template)
        structure = template.structure if template and template.structure else {}
        title_style = structure.get("titleStyle") or "Title"
        heading_style = structure.get("headingStyle") or "Heading 1"
        body_style = structure.get("bodyStyle") or "Normal"
        table_style = structure.get("tableStyle") or "Table Grid"

        title = document.add_paragraph(style=self._existing_style(document, title_style, "Title"))
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(report.report_name)
        title_run.bold = True
        title_run.font.size = Pt(18)

        body_style_name = self._existing_style(document, body_style, "Normal")
        document.add_paragraph(f"报告类型：{report_type_label(report.report_type)}", style=body_style_name)
        document.add_paragraph(f"主题：{report.topic or ''}", style=body_style_name)
        document.add_paragraph(f"电厂：{report.plant or ''}", style=body_style_name)
        document.add_paragraph(f"专业：{report.major or ''}", style=body_style_name)
        document.add_paragraph(f"年份：{report.year or ''}", style=body_style_name)

        for chapter in outline:
            level = min(max(chapter.level, 1), 4)
            if level == 1:
                heading = document.add_paragraph(style=self._existing_style(document, heading_style, "Heading 1"))
                heading.add_run(f"{chapter.chapter_no} {chapter.title}")
            else:
                document.add_heading(f"{chapter.chapter_no} {chapter.title}", level=level)
            content = contents_by_chapter.get(chapter.id)
            if content and content.content:
                document.add_paragraph(content.content, style=body_style_name)
            else:
                document.add_paragraph("本章节暂无保存内容。", style=body_style_name)
            for table_data in (content.tables if content else []) or []:
                document.add_paragraph(table_data.get("title", "表格"), style=body_style_name)
                headers = table_data.get("headers", [])
                rows = table_data.get("rows", [])
                table = document.add_table(rows=1, cols=max(len(headers), 1))
                table.style = self._existing_table_style(document, table_style, "Table Grid")
                for index, header in enumerate(headers or ["内容"]):
                    table.rows[0].cells[index].text = str(header)
                for row_data in rows:
                    row = table.add_row()
                    for index, value in enumerate(row_data[: len(table.columns)]):
                        row.cells[index].text = str(value)

        document.save(path)

    def _new_document(self, template: ReportTemplate | None) -> DocxDocument:
        template_path = Path(template.file_path) if template and template.file_path else None
        if template_path and template_path.exists() and template_path.suffix.lower() == ".docx":
            document = Document(template_path)
            self._clear_body(document)
            return document
        return Document()

    def _clear_body(self, document: DocxDocument) -> None:
        body = document.element.body
        for child in list(body):
            if child.tag.endswith("}sectPr"):
                continue
            body.remove(child)

    def _existing_style(self, document: DocxDocument, preferred: Any, fallback: str) -> str:
        style_name = str(preferred or fallback)
        try:
            document.styles[style_name]
            return style_name
        except KeyError:
            return fallback

    def _existing_table_style(self, document: DocxDocument, preferred: Any, fallback: str) -> str:
        style_name = str(preferred or fallback)
        try:
            document.styles[style_name]
            return style_name
        except KeyError:
            return fallback