from pathlib import Path

from docx import Document
from sqlalchemy.orm import Session

from app.entity.material import Material, MaterialTag

MAX_MATERIALS = 5
MAX_CHARS_PER_MATERIAL = 1000
MAX_CONTEXT_CHARS = 3600


def build_material_context(
    db: Session,
    major: str | None,
    material_ids: list[int] | None = None,
) -> str:
    materials = _selected_materials(db, material_ids, major) if material_ids else _active_materials(db, major)
    blocks: list[str] = []
    for item in materials[:MAX_MATERIALS]:
        major_tag = _tag(db, item.id, "major") or "未指定专业"
        text = _extract_material_text(item)
        if not text:
            text = item.description or "该素材暂无法提取正文，可参考素材名称、类型和描述。"
        text = _compact(text)[:MAX_CHARS_PER_MATERIAL]
        blocks.append(
            f"素材名称：{item.name}\n"
            f"素材类型：{item.type}\n"
            f"适用专业：{major_tag}\n"
            f"素材摘要：{text}"
        )
    return "\n\n".join(blocks)[:MAX_CONTEXT_CHARS]


def _selected_materials(db: Session, material_ids: list[int] | None, major: str | None) -> list[Material]:
    ids = [item for item in material_ids or [] if item]
    if not ids:
        return []
    active_ids = {
        row.material_id
        for row in db.query(MaterialTag)
        .filter(MaterialTag.tag_key == "status", MaterialTag.tag_value == "enabled")
        .all()
    }
    items = db.query(Material).filter(Material.id.in_(ids)).all()
    item_by_id = {
        item.id: item
        for item in items
        if item.id in active_ids and _material_major_allowed(_tag(db, item.id, "major"), major)
    }
    return [item_by_id[item_id] for item_id in ids if item_id in item_by_id]


def _active_materials(db: Session, major: str | None) -> list[Material]:
    active_ids = [
        row.material_id
        for row in db.query(MaterialTag)
        .filter(MaterialTag.tag_key == "status", MaterialTag.tag_value == "enabled")
        .all()
    ]
    query = db.query(Material).filter(Material.id.in_(active_ids or [-1]))
    if major:
        major_ids = [
            row.material_id
            for row in db.query(MaterialTag)
            .filter(
                MaterialTag.tag_key == "major",
                MaterialTag.tag_value.in_([major, "", "综合", "通用", "general", "common"]),
            )
            .all()
        ]
        if major_ids:
            query = query.filter(Material.id.in_(major_ids))
    return query.order_by(Material.id.desc()).all()


def _material_major_allowed(material_major: str | None, report_major: str | None) -> bool:
    if not report_major:
        return True
    normalized = (material_major or "").strip()
    return normalized in {"", report_major, "综合", "通用", "general", "common"}


def _tag(db: Session, material_id: int, key: str) -> str | None:
    tag = (
        db.query(MaterialTag)
        .filter(MaterialTag.material_id == material_id, MaterialTag.tag_key == key)
        .first()
    )
    return tag.tag_value if tag else None


def _extract_material_text(item: Material) -> str:
    path = Path(item.file_path or "")
    if not path.exists() or not path.is_file():
        return ""
    suffix = path.suffix.lower()
    try:
        if suffix in {".txt", ".md", ".csv"}:
            return _read_text_file(path)
        if suffix == ".docx":
            document = Document(path)
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            table_text: list[str] = []
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        table_text.append(" | ".join(cells))
            return "\n".join(paragraphs + table_text)
    except Exception:
        return ""
    return ""


def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_bytes().decode("utf-8", errors="ignore")


def _compact(text: str) -> str:
    return " ".join(text.split())
